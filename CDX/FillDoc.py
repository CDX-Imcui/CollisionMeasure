import time

from docx import Document
from docx.shared import Inches
from PIL import Image
from datetime import datetime
import os


def get_time_str(delta=0, fmt="%Y-%m-%d %H:%M:%S"):
    """
    :return: 2024-08-31 10:41:45
    """
    my_time = int(time.time()) + delta * 24 * 3600
    return time.strftime(fmt, time.localtime(my_time))

from docx.shared import Inches, Pt

def get_cell_dimensions(cell):
    # 获取列宽（以EMU为单位，1英寸 = 914400 EMU）
    try:
        width_emu = cell._tc.width
        width_inch = width_emu / 914400
    except:
        width_inch = 5.0  # 默认宽度

    # 行高需通过所在行的高度属性推测，通常设定在 trHeight，单位为 1/20 Point
    try:
        tr = cell._tc.getparent()
        height_twips = tr.trPr.trHeight.val
        # 合理范围判断：小于10英寸（14400 twips/inch * 10）
        if height_twips and height_twips < 144000:
            height_inch = height_twips / 1440  # 1英寸 = 1440 twips
        else:
            height_inch = 5.0
    except:
        height_inch = 5.0

    return width_inch, height_inch

def fill_word_template(template_path, output_path, fill_data):
    """
    在 Word 模板中的表格指定单元格填充文字和图像。

    :param template_path: str, 模板文件路径
    :param output_path: str, 填写后保存文件路径
    :param fill_data: list of dict, 每个 dict 形式为：
                      {
                          "table_index": 表格索引（从0开始），
                          "row": 行号（从0开始），
                          "col": 列号（从0开始），
                          "text": 填写的文字，可选，
                          "image_path": 图像路径，可选
                      }
    """
    doc = Document(template_path)

    for item in fill_data:
        table_index = item["table_index"]
        row = item["row"]
        col = item["col"]
        text = item.get("text", "")
        image_path = item.get("image_path", None)

        table = doc.tables[table_index]
        cell = table.cell(row, col)

        # 清空原有内容
        cell.text = ""

        if text:
            cell.text = text

        if image_path and os.path.exists(image_path):
            try:
                # 添加图片，需新建段落插入图像
                paragraph = cell.add_paragraph()
                paragraph.alignment = 1  # 居中对齐
                run = paragraph.add_run()

                with Image.open(image_path) as img:
                    img_width, img_height = img.size
                    img_ratio = img_height / img_width

                cell = table.cell(2, 0)
                cell_width, cell_height = get_cell_dimensions(cell)
                cell_ratio = cell_height / cell_width

                if img_ratio > cell_ratio:
                    target_height = cell_height
                    target_width = target_height / img_ratio
                else:
                    target_width = cell_width
                    target_height = target_width * img_ratio

                run.add_picture(image_path, width=Inches(target_width), height=Inches(target_height))
            except Exception as e:
                print(f"图片插入失败: {e}")
    try:
        doc.save(output_path)
    except PermissionError:
        print(f"保存失败：文件 {output_path} 正在使用中，请关闭后重试。")



if __name__ == "__main__":
    # 获取当前时间并格式化为指定格式
    input_path = "事故报告_模板.docx"
    output_path = "事故报告.docx"
    current_time = datetime.now().strftime("%Y年%m月%d日%H时%M分%S秒")
    address = "浙江省杭州市西湖区教工路"
    image_path = "E:\\22.jpg"
    fill_data = [
        {"table_index": 0, "row": 0, "col": 1, "text": current_time},
        {"table_index": 0, "row": 1, "col": 1, "text": address},
        {"table_index": 0, "row": 2, "col": 0, "image_path": image_path}
    ]

    fill_word_template(input_path, output_path, fill_data)
