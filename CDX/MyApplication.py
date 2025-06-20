# -*- coding: utf-8 -*-
import os
import shutil
import sys

import cv2
import numpy as np
import vtk
from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtCore import QSettings, pyqtSignal, QObject
from PyQt5.QtWidgets import QFrame, QLabel, QVBoxLayout, QPushButton, QInputDialog, QHBoxLayout
from matplotlib import pyplot as plt
from vtkmodules.qt.QVTKRenderWindowInteractor import QVTKRenderWindowInteractor

from Clickable_Label import Clickable_Label
from from_matrix import DepthBackProjector
from mainwindow import Ui_MainWindow
from PointCloud import showPointCloud, getPointCloud, mark_point, draw_line_between_points
from Threads import AddImagesWorker, Parsing_video, ColmapWorker, Align_according_to_LicensePlate_Worker, \
    SemanticSegmentation_Worker
from datetime import datetime
from glob import glob
from FillDoc import fill_word_template
import matplotlib as mpl

# 清理临时文件
def cleardir_ine(dirname):
    if not os.path.exists(dirname):
        os.makedirs(dirname)
        return

    filelist = []
    filelist = os.listdir(dirname)
    for f in filelist:
        filepath = os.path.join(dirname, f)
        if os.path.isfile(filepath):
            os.remove(filepath)
        elif os.path.isdir(filepath):
            shutil.rmtree(filepath, True)


class MyStyle(vtk.vtkInteractorStyleTrackballCamera):
    def __init__(self, callback=None, callback_right=None):
        super().__init__()
        self.AddObserver("KeyPressEvent", self.on_key_press)
        self.AddObserver("KeyReleaseEvent", self.on_key_release)
        self.AddObserver("LeftButtonPressEvent", self.on_left_click)
        self.AddObserver("RightButtonPressEvent", self.on_right_click)
        self.space_pressed = False
        self.callback = callback
        self.callback_right = callback_right

    def on_key_press(self, obj, event):
        key = self.GetInteractor().GetKeySym()
        if key == "space":
            self.space_pressed = True
            # print("Space pressed")

    def on_key_release(self, obj, event):
        key = self.GetInteractor().GetKeySym()
        if key == "space":
            self.space_pressed = False
            # print("Space released")

    def on_left_click(self, obj, event):
        if self.space_pressed:
            Position = self.GetInteractor().GetEventPosition()
            self.callback(Position)
        else:
            # 保留默认旋转行为
            self.OnLeftButtonDown()

    def on_right_click(self, obj, event):
        Position = self.GetInteractor().GetEventPosition()
        self.callback_right(clickPos=Position)


class MainWindow(QtWidgets.QMainWindow, Ui_MainWindow):
    def __init__(self):
        super(MainWindow, self).__init__()
        self.setupUi(self)
        self.add_ui()
        self.ColmapThread = None
        self.SplatThread = None
        self.image_paths = []
        self.LicensePlate = 0.44  # 米 （）
        self.Unit_distance_length = None  # （）

        # 初始化时设置临时工作目录
        self.WORK_DIR = os.path.join(os.getcwd(), 'WORK_DIR')
        self.save_path = os.path.join(self.WORK_DIR, "save_images")
        # 确保保存目录存在
        os.makedirs(os.path.join(os.getcwd(), 'WORK_DIR',"save_images"), exist_ok=True)
        if not os.path.exists(self.WORK_DIR):
            os.makedirs(self.WORK_DIR)
        else:
            cleardir_ine(os.path.join(self.WORK_DIR, "distorted"))
            cleardir_ine(os.path.join(self.WORK_DIR, "images"))
            cleardir_ine(os.path.join(self.WORK_DIR, "input"))
            cleardir_ine(os.path.join(self.WORK_DIR, "sparse"))
            cleardir_ine(os.path.join(self.WORK_DIR, "stereo"))
            cleardir_ine(os.path.join(self.WORK_DIR, "save_images"))

        os.makedirs(self.save_path, exist_ok=True)
        self.best_image_id = None  # int （）
        self.best_image_name = None  # 'image_00059.jpg'（）
        self.best_points = None
        self.have_plane = False
        self.plane = None

        self.thumbnail_width = 200  # 设置缩略图宽度
        # self.sidebar_visible = False  # 初始化侧边栏可见状态
        self.settings = QSettings('CUIDONGXU', 'point_cloud_splat')  # 指定组织名和应用名时，自动确定数据的存储位置和格式
        # 从设置中加载最近文件列表
        self.recent_files_list = self.settings.value('recentFiles', [])  # 通过 setValue 和 value 方法可以写入和读取具体的设置项

        self.point_cloud = vtk.vtkPolyData()  # 用于存储加载的点云数据
        self.auxiliary_actors = []  # 存储除了点云之外自己加的标注、线段、标签等所有 actor
        # 测距功能所需的变量
        self.measuring = False  # 是否处于测距状态
        self.selected_points = []  # 存储已选择的点
        self.distance_labels = []  # 存储显示距离的标签
        self.point_labels = []  # 存储显示点坐标的标签

        # 多组点对测量相关变量
        self.point_pairs = []  # 存储所有点对 [(p1, p2), (p3, p4), ...]
        self.distances = []  # 存储所有点对的计算距离
        self.real_distances = []  # 存储实际距离，初始为None
        self.reference_pair_index = -1  # 参考点对的索引，-1表示未设置

        # 连接信号和槽
        self.actionOpenPointCloud.triggered.connect(self.openPointCloud)
        self.actionRecentFiles.triggered.connect(self.show_recent_files)
        self.actionExit.triggered.connect(self.close)
        self.actionReconstruct.triggered.connect(self.startColmap)
        # self.mesure_distance.triggered.connect(self.start_measure_distance) #暂时先不做
        self.export_docx_Action.triggered.connect(self.export_docx)
        # 初始化最近文件菜单
        self.updateRecentFilesMenu()
        self.import_video.triggered.connect(self.addVideo)
        self.import_images.triggered.connect(self.addImageFolder)

    def add_ui(self):
        self.setWindowTitle("测距")
        self.actionReconstruct = QtWidgets.QAction("开始构建", self)
        self.menubar.addAction(self.actionReconstruct)
        self.mesure_distance = QtWidgets.QAction(" ", self)  # 开始测距
        self.menubar.addAction(self.mesure_distance)
        self.export_docx_Action = QtWidgets.QAction("生成报告", self)
        self.menubar.addAction(self.export_docx_Action)

        self.centralwidget = QtWidgets.QWidget()
        self.centralwidget.setObjectName("centralwidget")
        self.setCentralWidget(self.centralwidget)

        # 创建标签页控件
        self.tabWidget = QtWidgets.QTabWidget(self.centralwidget)
        self.tabWidget.setObjectName("tabWidget")
        # 使用布局管理器自动重新排列和调整大小来适应新的窗口尺寸!否则铺不满
        self.verticalLayout = QtWidgets.QVBoxLayout(self.centralwidget)
        self.verticalLayout.setObjectName("verticalLayout")
        self.verticalLayout.addWidget(self.tabWidget)  # 添加 tabWidget 到布局中

        # 标签页0 - 缩略图
        self.tabThumbnails = QtWidgets.QWidget()
        self.tabThumbnails.setObjectName("tabThumbnails")
        self.scrollAreaThumbnails = QtWidgets.QScrollArea(self.tabThumbnails)
        self.scrollAreaThumbnails.setWidgetResizable(True)
        self.scrollAreaContentThumbnails = QtWidgets.QWidget()
        self.scrollAreaContentThumbnails.setGeometry(QtCore.QRect(0, 0, 800, 600))  # 设置足够大的初始区域
        # QGridLayout网格布局，有序排列缩略图，支持动态添加或删除缩略图，并自动调整布局
        self.gridLayoutThumbnails = QtWidgets.QGridLayout(self.scrollAreaContentThumbnails)
        self.scrollAreaThumbnails.setWidget(self.scrollAreaContentThumbnails)
        self.verticalLayoutThumbnails = QtWidgets.QVBoxLayout(self.tabThumbnails)
        self.verticalLayoutThumbnails.addWidget(self.scrollAreaThumbnails)
        self.tabWidget.addTab(self.tabThumbnails, "缩略图")

        # 标签页1 - 图像
        self.tabImages = QtWidgets.QWidget()
        self.tabImages.setObjectName("tabImages")
        self.horizontalLayout = QtWidgets.QHBoxLayout(self.tabImages)
        self.scrollAreaImages = QtWidgets.QScrollArea(self.tabImages)
        self.scrollAreaImages.setWidgetResizable(True)
        self.scrollAreaImages.setObjectName("scrollAreaImages")
        self.scrollAreaContentImages = QtWidgets.QWidget()
        self.scrollAreaContentImages.setObjectName("scrollAreaContentImages")
        self.horizontalLayoutImages = QtWidgets.QHBoxLayout(self.scrollAreaContentImages)  # 水平布局管理器
        self.scrollAreaContentImages.setLayout(self.horizontalLayoutImages)
        self.scrollAreaImages.setWidget(self.scrollAreaContentImages)
        self.horizontalLayout.addWidget(self.scrollAreaImages)
        self.tabWidget.addTab(self.tabImages, "图像")

        # 标签页2 - 点云显示
        self.tabPointCloud = QtWidgets.QWidget()
        self.tabPointCloud.setObjectName("tabPointCloud")
        # self.tabWidget.addTab(self.tabPointCloud, "点云")
        self.tabWidget.addTab(self.tabPointCloud, "null")  # 暂时先这样
        # 设置水平布局
        self.main_horizontalLayout = QtWidgets.QHBoxLayout(self.tabPointCloud)
        # 添加一个QVTKRenderWindowInteractor作为点云显示区域
        self.vtkWidget = QVTKRenderWindowInteractor(self.tabPointCloud)
        self.main_horizontalLayout.addWidget(self.vtkWidget, 10)
        self.renderer = vtk.vtkRenderer()
        self.renderer.SetBackground(0.15, 0.15, 0.16)
        self.vtkWidget.GetRenderWindow().AddRenderer(self.renderer)
        self.vtkWidget.GetRenderWindow().Render()
        # 拖动点云的操作方式
        # style = vtkInteractorStyleTrackballCamera()
        style = MyStyle(callback=self.align_center_click, callback_right=self.pick_callback)
        self.vtkWidget.GetRenderWindow().GetInteractor().SetInteractorStyle(style)
        self.vtkWidget.Initialize()  # 初始化 vtkWidget
        # 存储渲染器和交互器的引用，以便在后续操作中使用
        self.renderWindowInteractor = self.vtkWidget.GetRenderWindow().GetInteractor()
        # 显示控制台信息
        self.textEditConsole = QtWidgets.QTextEdit(self.tabPointCloud)
        self.textEditConsole.setObjectName("textEditConsole")
        self.textEditConsole.setReadOnly(True)  # 设置 QTextEdit 为只读
        self.main_horizontalLayout.addWidget(self.textEditConsole, 2)  # 伸展因子1
        # 确保布局填满整个tab
        self.tabPointCloud.setLayout(self.main_horizontalLayout)

        # 右边侧边栏（可隐藏）
        self.right_sidebar = QFrame()
        self.main_horizontalLayout.addWidget(self.right_sidebar, 3)
        self.right_sidebar.setStyleSheet("background-color: #DDEEFF;")
        self.right_sidebar_layout = QVBoxLayout(self.right_sidebar)

        # 测距功能的UI组件
        sidebar_title = QLabel("点云测距")
        # sidebar_title.setStyleSheet("font-weight: bold; font-size: 14px;")
        self.right_sidebar_layout.addWidget(sidebar_title)
        # 添加重置按钮
        self.reset_button = QPushButton("清除所有点")
        self.reset_button.clicked.connect(self.reset_measuring)
        self.right_sidebar_layout.addWidget(self.reset_button)
        # 添加点坐标信息区域
        self.points_info_label = QLabel("点坐标信息:")
        self.points_info_label.setStyleSheet("font-weight: bold;")
        self.right_sidebar_layout.addWidget(self.points_info_label)
        # 距离信息区域
        self.distance_info_label = QLabel("距离信息:")
        self.distance_info_label.setStyleSheet("font-weight: bold;")
        self.right_sidebar_layout.addWidget(self.distance_info_label)

        self.right_sidebar_layout.addStretch()  # 添加一个“弹性空间”
        self.right_sidebar.hide()  # 启动时隐藏侧边栏

        self.actionRecentFiles.setMenu(QtWidgets.QMenu())

    def retranslateUi(self, MainWindow):
        super().retranslateUi(MainWindow)
        _translate = QtCore.QCoreApplication.translate  # 用于自动连接信号和槽
        MainWindow.setWindowTitle(_translate("MainWindow", "图像密集匹配系统"))
        # MainWindow.setWindowIcon(QtGui.QIcon('cdx.png'))
        MainWindow.setWindowIcon(QtGui.QIcon('_internal/measure.ico'))

    def openPointCloud(self, file_path=None):
        # 对 file_path 要判断，从对话框选择打开是空的，需要文件对话框���从最���是指定了路径，需要跳过文件对话框
        if not file_path:
            options = QtWidgets.QFileDialog.Options()
            file_path, _ = QtWidgets.QFileDialog.getOpenFileName(self,
                                                                 "选择点云文件", "",
                                                                 "点云文件(*.ply *.obj *.stl *.vtk)",
                                                                 options=options)
        if file_path:
            # 检查文件是否存在  *.ply *.obj *.stl *.vtk
            if not os.path.exists(file_path):
                QtWidgets.QMessageBox.warning(self, "文件打开错误", "文件不存在或已被删除: " + file_path)
                # 更新列表，移除不存在的文件
                if file_path in self.recent_files_list:
                    self.recent_files_list.remove(file_path)
                    self.updateRecentFilesMenu()
                return

            try:
                actors = self.renderer.GetActors()
                actors.InitTraversal()
                for i in range(actors.GetNumberOfItems()):
                    actor = actors.GetNextActor()
                    self.renderer.RemoveActor(actor)

                self.point_cloud = getPointCloud(file_path)
                if self.point_cloud.GetNumberOfPoints() == 0:
                    print("点云数据为空")
                    return
                showPointCloud(self.point_cloud, self.vtkWidget, self.renderer)  # 显示点云

                # 点云文件成功打开后，更新最近文件列表和菜单。   工作路径中的文件会被清除，���应该保存到列表中
                if file_path not in self.recent_files_list:  # 避免最近文件重复
                    self.recent_files_list.append(file_path)
                    self.recent_files_list = self.recent_files_list[-10:]  # 保持列表只有最后10个文件
                    self.updateRecentFilesMenu()  # 更新最近文件菜单
            except Exception as e:
                print(f"读取或显示点云时出现错误：{e}")
        else:
            print("未选择文件")

    def show_recent_files(self):
        # 实现显示最近文件的功能
        print("显示最近文件")
        # 这里添加显示最近文件的代码

    def start_measure_distance(self):
        """启用或禁用测距功能"""
        if not self.measuring:
            # 打开测距模式
            self.right_sidebar.show()
            self.mesure_distance.setText("关闭测距")
            self.measuring = True

            # 添加点击事件监听器
            # self.observer_id = self.renderWindowInteractor.AddObserver("RightButtonPressEvent", self.pick_callback)
            print("测距模式已启用，请点击点云上的点")
        else:
            # 关闭测距模式
            self.right_sidebar.hide()
            self.mesure_distance.setText("开始测距")
            self.measuring = False
            # # 移除点击事件监听器
            # if hasattr(self, 'observer_id'):
            #     self.renderWindowInteractor.RemoveObserver(self.observer_id)
            #     delattr(self, 'observer_id')
            print("测距模式已关闭")

    def pick_callback(self, clickPos):
        if not self.measuring:
            return
        # clickPos = self.renderWindowInteractor.GetEventPosition()
        picker = vtk.vtkPointPicker()
        # 设置点选取的容差（像素单位）
        picker.SetTolerance(3)
        # 执行选点操作
        picker.Pick(clickPos[0], clickPos[1], 0, self.renderer)
        pid = picker.GetPointId()

        if pid != -1:
            pos = self.point_cloud.GetPoint(pid)
            pos_array = np.array(pos)
            self.selected_points.append(pos_array)
            print(f"已选中点: {pos}")
            # 标记选中的点
            mark_point(pos, self.vtkWidget, self.renderer, self.auxiliary_actors)

            # 当选择了两个点时，进行距离测量
            if len(self.selected_points) == 2:
                p1, p2 = self.selected_points
                dist = np.linalg.norm(p1 - p2)
                print(f"两点距离: {dist:.4f}")

                # 存储点对和距离
                self.point_pairs.append((p1.copy(), p2.copy()))
                self.distances.append(dist)
                self.real_distances.append(None)  # 初始时实际距离未知
                pair_index = len(self.point_pairs) - 1
                # 在3D场景中绘制连线
                draw_line_between_points(p1, p2, pair_index, self.distances, self.real_distances, self.vtkWidget,
                                         self.renderer,
                                         self.auxiliary_actors)
                # 在侧边栏中更新距离信息
                self.update_sidebar_distance_info()
                # 重置选中点，准备选择下一对点
                self.selected_points = []

    def input_real_distance(self, index):
        """用户输入参考点对的实际距离"""
        dist, ok = QInputDialog.getDouble(self, "输入实际距离", f"为点对 {index + 1} 输入实际长度:", decimals=4)
        if ok:
            self.set_as_reference_pair(index, dist)

    def set_as_reference_pair(self, index, real_length=None):
        """设置某点对为参考点对，并根据比例更新其他所有实际距离"""
        ref_measured = self.distances[index]
        if ref_measured == 0:
            QtWidgets.QMessageBox.warning(self, "无效参考", "参考点对的测量距离为零，无法计算比例。")
            return

        if real_length is None:
            # 默认按原来的比例换算
            real_length = self.real_distances[index]

        if real_length is None:
            # 如果没有输入实际距离，也无法计算
            QtWidgets.QMessageBox.warning(self, "实际距离缺失", "请先为该点对输入实际距离。")
            return

        # 设置参考点对
        self.reference_pair_index = index
        self.real_distances[index] = real_length

        # 计算比例
        scale_ratio = real_length / ref_measured

        # 应用比例到其他点对
        for i, d in enumerate(self.distances):
            if i == index:
                continue
            self.real_distances[i] = d * scale_ratio

        self.update_sidebar_distance_info()
        self.vtkWidget.GetRenderWindow().Render()

    def update_sidebar_distance_info(self):
        """更新侧边栏中显示的距离信息"""
        # 清空现有距离信息
        # 移除之前添加的距离标签
        for label in self.distance_labels:
            self.right_sidebar_layout.removeWidget(label)
            label.deleteLater()
        self.distance_labels = []

        # 添加新的距离信息
        for i, (dist, real_dist) in enumerate(zip(self.distances, self.real_distances)):
            # 创建距离信息文本
            if real_dist is not None:
                text = f"点对 {i + 1}: 测量距离 = {dist:.2f}, 实际距离 = {real_dist:.2f}"
            else:
                text = f"点对 {i + 1}: 测量距离 = {dist:.2f}"

            # 如果是参考点对，添加标记
            if i == self.reference_pair_index:
                text += " [参考点对]"

            label = QLabel(text)
            if i == self.reference_pair_index:
                label.setStyleSheet("color: red; font-weight: bold;")
            else:
                label.setStyleSheet("color: blue;")

            # 添加设置为参考点对的按钮
            container = QFrame()
            container_layout = QHBoxLayout(container)
            container_layout.setContentsMargins(0, 0, 0, 0)

            container_layout.addWidget(label)

            if i != self.reference_pair_index:
                set_ref_button = QPushButton("设为参考")
                set_ref_button.setMaximumWidth(80)
                set_ref_button.clicked.connect(lambda checked, idx=i: self.set_as_reference_pair(idx))
                container_layout.addWidget(set_ref_button)

            # 添加输入实际距离的按钮
            input_button = QPushButton("输入实际距离")
            input_button.setMaximumWidth(100)
            input_button.clicked.connect(lambda checked, idx=i: self.input_real_distance(idx))
            container_layout.addWidget(input_button)

            # 将整个容器添加到侧边栏
            self.right_sidebar_layout.insertWidget(4 + i, container)
            self.distance_labels.append(container)

        # 重新渲染场景
        self.vtkWidget.GetRenderWindow().Render()

    def set_as_reference_pair(self, index):
        """设置指定索引的点对为参考点对"""
        if 0 <= index < len(self.point_pairs):
            # 更新参考点对索引
            self.reference_pair_index = index
            print(f"已将点对 {index + 1} 设置为参考点对")

            # 如果参考点对有实际距离，则更新所有其他点对的实际距离
            if self.real_distances[index] is not None:
                self.update_all_real_distances()

            # 更新侧边栏信息
            self.update_sidebar_distance_info()

    def input_real_distance(self, index):
        """输入指定点对的实际距离"""
        if 0 <= index < len(self.point_pairs):
            # 提示用户输入实际距离
            real_dist, ok = QtWidgets.QInputDialog.getDouble(
                self,
                f"输入点对 {index + 1} 的实际距离",
                "请输入实际距离值:",
                value=0.0,
                min=0.0,
                max=1000000.0,
                decimals=4
            )

            if ok:
                # 更新实际距离
                self.real_distances[index] = real_dist
                print(f"点对 {index + 1} 的实际距离设置为 {real_dist:.4f}")

                # 如果当前点对是参考点对，则更新所有其他点对的实际距离
                if index == self.reference_pair_index:
                    self.update_all_real_distances()

                # 更新3D场景中的距离标签
                p1, p2 = self.point_pairs[index]
                self.update_distance_text(index)

                # 更新侧边栏信息
                self.update_sidebar_distance_info()

    def update_all_real_distances(self):
        """根据参考点对的比例更新所有点对的实际距离"""
        if self.reference_pair_index == -1:
            return

        ref_idx = self.reference_pair_index
        ref_measured = self.distances[ref_idx]
        ref_real = self.real_distances[ref_idx]

        if ref_measured <= 0 or ref_real is None:
            return

        # 计算比例
        scale_factor = ref_real / ref_measured

        # 更新所有点对的实际距离
        for i, dist in enumerate(self.distances):
            if i != ref_idx:  # 跳过参考点对
                self.real_distances[i] = dist * scale_factor
                print(f"点对 {i + 1} 的实际距离更新为 {self.real_distances[i]:.4f}")

        # 更新所有3D场景中的距离标签
        for index in range(len(self.point_pairs)):
            self.update_distance_text(index)

    def update_distance_text(self, index):
        """更新3D场景中指定点对的距离标签"""
        # 这里需要重新创建或修改3D场景中的文本
        # 由于VTK的限制，我们需要重新绘制连线和标签
        p1, p2 = self.point_pairs[index]
        draw_line_between_points(p1, p2, index, self.distances, self.real_distances, self.vtkWidget, self.renderer,
                                 self.auxiliary_actors)

    def reset_measuring(self):
        """重置测距工具，清除所有已选择的点和测量信息"""
        # 清除所有点对和距离数据
        self.point_pairs = []
        self.distances = []
        self.real_distances = []
        self.reference_pair_index = -1
        self.selected_points = []

        # 清除所有显示的点和线
        for actor in self.auxiliary_actors:
            self.renderer.RemoveActor(actor)
        self.auxiliary_actors.clear()

        # 清除右侧边栏中的点坐标和距离信息
        for label in self.distance_labels:
            self.right_sidebar_layout.removeWidget(label)
            label.deleteLater()
        self.distance_labels = []
        # 清除右侧边栏中的点坐标和距离信息
        self.points_info_label.setText("点坐标信息:")
        self.distance_info_label.setText("距离信息:")

    def updateRecentFilesMenu(self):
        # 检查是否已经有一个名为 actionRecentFiles 的属性，如果没有，就创建一个新的 QAction
        if not hasattr(self, 'actionRecentFiles'):
            self.actionRecentFiles = QtWidgets.QAction("最近保存文件", self)
            self.menuFileActions.addAction(self.actionRecentFiles)
            recentFilesMenu = QtWidgets.QMenu()
            self.actionRecentFiles.setMenu(recentFilesMenu)
        else:
            self.actionRecentFiles.menu().clear()

        recentFilesMenu = self.actionRecentFiles.menu()

        if self.recent_files_list:
            for file_path in reversed(self.recent_files_list[-10:]):
                action = QtWidgets.QAction(os.path.basename(file_path), self)
                # 绑定动作的触发事件到打开文件的槽函数，使用lambda确保传递正确的file_path
                action.triggered.connect(lambda checked, path=file_path: self.open_point_cloud(path))  # 避免延迟绑定
                recentFilesMenu.addAction(action)

            # 当存在记录时才添加分隔符和清除记录选项
            recentFilesMenu.addSeparator()
            clearAction = QtWidgets.QAction("清除记录", self)
            clearAction.triggered.connect(self.clearRecentFiles)
            recentFilesMenu.addAction(clearAction)

    def clearRecentFiles(self):
        self.recent_files_list = []  # 清空最近文件列表
        self.updateRecentFilesMenu()  # 重新更新菜单
        self.settings.setValue('recentFiles', self.recent_files_list)  # 更新设置，以便在下次启动时保持清空状态

    def update_point_info(self):
        """更新界面上显示的点坐标信息"""
        info_text = "点坐标信息:\n"

        for i, point in enumerate(self.picked_points):
            # 格式化显示坐标，��留3位小数
            point_text = f"点{i + 1}: ({point[0]:.3f}, {point[1]:.3f}, {point[2]:.3f})"
            info_text += point_text + "\n"

            # 创建或更新标签
            if i < len(self.point_labels):
                self.point_labels[i].setText(point_text)
            else:
                label = QLabel(point_text)
                self.right_sidebar_layout.insertWidget(3 + i, label)  # 在"点坐标信息"标签后插入
                self.point_labels.append(label)

        # 更新总信息标签
        self.points_info_label.setText(info_text)

    def align_center_click(self, click_pos):
        """空格加单击 将点击的点设置为视图中心"""
        # print("双击事件触发")
        picker = vtk.vtkPointPicker()
        picker.Pick(click_pos[0], click_pos[1], 0, self.renderer)
        point_id = picker.GetPointId()

        if point_id != -1:
            pos = self.point_cloud.GetPoint(point_id)
            print(f"双击点坐标: {pos}")
            camera = self.renderer.GetActiveCamera()

            # 获取当前相机位置和焦点
            old_pos = np.array(camera.GetPosition())
            old_focal = np.array(camera.GetFocalPoint())

            # 将新的焦点设置为点击点
            new_focal = np.array(pos)
            camera.SetFocalPoint(*new_focal)

            # 保持相机相对于焦点的方向和距离不变，平移位置
            direction = old_pos - old_focal
            new_pos = new_focal + direction
            camera.SetPosition(*new_pos)

            # 重新渲染
            self.vtkWidget.GetRenderWindow().Render()
        else:
            print("未拾取到点")

    def addVideo(self):
        if not hasattr(self, 'Parsing_video_worker') or not self.Parsing_video_worker.isRunning():
            cleardir_ine(self.WORK_DIR)
            video_path, _ = QtWidgets.QFileDialog.getOpenFileName(self, '选择视频', '',
                                                                  '视频文件 (*.mp4)')
            if video_path:
                self.Parsing_video_worker = Parsing_video(video_path, self.WORK_DIR, self.image_paths)
                self.Parsing_video_worker.finished.connect(self.showImages)
                self.Parsing_video_worker.start()
        else:
            QtWidgets.QMessageBox.warning(self, "警告", "正在解析视频")

    def addImageFolder(self):
        if not hasattr(self, 'add_images_worker') or not self.add_images_worker.isRunning():
            cleardir_ine(self.WORK_DIR)
            folder = QtWidgets.QFileDialog.getExistingDirectory(self, '选择文件夹', '')
            if folder:
                # 获取文件夹内所有图像文件
                images = [os.path.join(folder, f) for f in os.listdir(folder)
                          if
                          os.path.isfile(os.path.join(folder, f)) and f.lower().endswith(
                              ('.png', '.jpg', '.heic', '.jpeg', '.bmp', '.tiff'))]
                self.add_images_worker = AddImagesWorker(images, self.WORK_DIR)
                self.add_images_worker.finished.connect(self.showImages)
                self.add_images_worker.start()
        else:
            QtWidgets.QMessageBox.warning(self, "警告", "正在添加图像")

    def showImages(self, file_paths):
        self.tabWidget.setCurrentIndex(0)  # 跳转到 缩略图标签页
        # 清空缩略图和图像标签页中所有内容
        self.clearLayout(self.gridLayoutThumbnails)
        self.clearLayout(self.horizontalLayoutImages)
        """同时更新 缩略图标签页、图像标签页"""
        for idx, file_path in enumerate(file_paths):
            # label = QtWidgets.QLabel()  # 每个缩略图是一个QLabel
            label = Clickable_Label(file_path)  # 创建一个可点击的标签
            label.clicked.connect(self.on_thumbnail_clicked_toPoint)
            # label.setAlignment(QtCore.Qt.AlignCenter)  # 设置标签内容居中
            # label.setStyleSheet("border: 1px solid #ccc;")  # 设置边框样式
            # label.setFixedSize(self.thumbnail_width, self.thumbnail_width)  #
            label.setObjectName(f"{os.path.basename(file_path)}")
            label.setToolTip(f"{os.path.basename(file_path)}")  # 设置鼠标悬停时显示的提示信息

            pixmap = QtGui.QPixmap(file_path)  # 缩小版的原始图像文件
            scaled_pixmap = pixmap.scaled(self.thumbnail_width, self.thumbnail_width, QtCore.Qt.KeepAspectRatio)
            label.setPixmap(scaled_pixmap)
            i = self.gridLayoutThumbnails.count()
            width = self.scrollAreaThumbnails.viewport().width()
            num_columns = max(1, width // (self.thumbnail_width + 10))  # 假设每个缩略图之间有10像素间隔
            row = i // num_columns
            column = i % num_columns
            self.gridLayoutThumbnails.addWidget(label, row, column)  # 将缩略图label添加到网格布局中
            # 每加载 10~20 张图片，就让 UI 刷新一次，避免界面假死
            if (idx + 1) % 20 == 0:
                QtCore.QCoreApplication.processEvents()

            # 图像标签页
            label_images = QtWidgets.QLabel()
            scaled_pixmap_large = pixmap.scaledToHeight(450, QtCore.Qt.SmoothTransformation)  # 将高度固定为450像素
            label_images.setPixmap(scaled_pixmap_large)
            self.horizontalLayoutImages.addWidget(label_images)
        self.updateGridLayout()
        width = sum([self.horizontalLayoutImages.itemAt(i).widget().width() for i in
                     range(self.horizontalLayoutImages.count())])
        self.scrollAreaContentImages.setMinimumWidth(width)
        self.scrollAreaContentImages.adjustSize()

    def on_thumbnail_clicked_toPoint(self, image_path):
        image_name = os.path.basename(image_path)
        base_name = os.path.splitext(image_name)[0]  # 获取不带扩展名的文件名

        PointCoordinate = []
        ChaCoordinate = []

        # cv2.imread 默认以 BGR 格式读取图像，将图片从 BGR 转换为 RGB，以便 matplotlib 正确显示颜色
        img = cv2.imread(image_path)
        img_rgb = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
        # 创建 Matplotlib 图形窗口，并设置响应更快的后端
        plt.switch_backend('Qt5Agg')  # 使用Qt5Agg图形后端

        fig = plt.figure(figsize=(10, 8))
        gs = fig.add_gridspec(2, 1, height_ratios=[9, 1])
        fig.subplots_adjust(top=0.85, bottom=0.05, left=0.05, right=0.95)  # 为顶部按钮预留空间

        # 创建主图形区域和按钮区域
        ax = fig.add_subplot(gs[0])  # 主图形区域
        button_ax = fig.add_subplot(gs[1])  # 按钮区域
        button_ax.axis('off')  # 隐藏按钮区域的坐标轴

        ax.imshow(img_rgb)
        ax.axis('off')  # 关闭坐标轴

        if self.Unit_distance_length is None:
            return

        projector = DepthBackProjector(self.WORK_DIR, self.plane)
        projector.load_data(image_name)

        def onclick(event):
            if event.xdata is not None and event.ydata is not None:
                # 获取点击的像素坐标 (x, y)
                # event.xdata 和 event.ydata 是浮点数，需要转换为整数
                # 对于图像，y 是行，x 是列
                x = int(event.xdata + 0.5)  # 添加0.5并取整以提高精度
                y = int(event.ydata + 0.5)

                if 0 <= y < img.shape[0] and 0 <= x < img.shape[1]:  # 确保坐标在图片范围内
                    # b, g, r = img[y, x]  # 获取该像素的颜色值
                    # print(f"点击坐标 (X={x}, Y={y}) 的像素颜色 (RGB): ({r}, {g}, {b})")

                    t, error = projector.pixel_to_world(x, y)  # t:x, y, z（未经投影的点）一维NumPy数组
                    if (isinstance(t, (int, float)) and t == -1) or (
                            isinstance(t, np.ndarray) and np.all(t == -1)) and error == -1:
                        self.textEditConsole.append(f"点 ({x}, {y}) 在深度图中未找到对应的三维坐标")
                        # QtWidgets.QMessageBox.information(self, "点击无效", f"点 ({x}, {y}) 无效")
                        return
                    ax.plot(x, y, 'r+', markersize=10)  # 在图像上显示十字标记
                    plt.draw()  # 立即更新绘图
                    ChaCoordinate.append([x, y])

                    # # 确保t始终是一维NumPy数组
                    # if not isinstance(t, np.ndarray):
                    #     t = np.array(t, dtype=np.float64)

                    # PointCoordinate.append(t)  # 将三维坐标添加到列表中
                    # print("三维坐标 (world):",
                    #       [np.round(point, 4).tolist() if hasattr(point, 'tolist') else point for point in
                    #        PointCoordinate])
                    # print("三维坐标 (world):", np.round(PointCoordinate, 4))
                    # PointCoordinate中每有两个点，就绘制一条线，并标上距离
                    if len(ChaCoordinate) % 2 == 0:
                        # 绘制线段
                        p1 = ChaCoordinate[-2]
                        p2 = ChaCoordinate[-1]

                        line, error = projector.compute_distance(p1, p2)  # 计算距离
                        real_distance = line * self.Unit_distance_length  # 恢复成现实单位

                        ax.plot([ChaCoordinate[-2][0], ChaCoordinate[-1][0]],
                                [ChaCoordinate[-2][1], ChaCoordinate[-1][1]], 'g-')
                        # 计算线段中点位置用于放置距离标签
                        mid_x = (ChaCoordinate[-2][0] + ChaCoordinate[-1][0]) / 2
                        mid_y = (ChaCoordinate[-2][1] + ChaCoordinate[-1][1]) / 2
                        # 在线段中点显示距离值
                        ax.text(mid_x, mid_y, f"{real_distance:.2f}", color='red', fontsize=10,
                                ha='center', va='center')

        cid = fig.canvas.mpl_connect('button_press_event', onclick)

        def clear_points():
            """清空选点"""
            nonlocal PointCoordinate, ChaCoordinate
            PointCoordinate = []
            ChaCoordinate = []
            # 重新读取并显示图片以清除所有标记
            ax.clear()
            img_rgb = cv2.cvtColor(cv2.imread(image_path), cv2.COLOR_BGR2RGB)
            ax.imshow(img_rgb)
            ax.axis('off')
            plt.draw()  # 立即更新图像

        def save_images():
            """保存图像和选点"""
            if not ChaCoordinate:
                print("没有选点，无法保存")
                return

            # 保存图像
            fig.savefig(os.path.join(self.save_path, f"{base_name}.jpg"), bbox_inches='tight', pad_inches=0.1, dpi=600)
            print(f"图像已保存到: {self.save_path}")
            QtWidgets.QMessageBox.information(None, "保存成功", f"图像已保存到: {self.save_path}")


        mpl.rcParams['keymap.save'] = []
        def on_key(event):
            if event.key == 'ctrl+c':  # 按下 C 键
                clear_points()
            elif event.key == 'ctrl+s':  # 按下 S 键
                save_images()
            elif event.key == 'q':  # 按下 Q 键退出
                plt.close(fig)
        fig.canvas.mpl_connect('key_press_event', on_key)

        # 分离交互区域，确保tight_layout不影响按钮位置
        fig.canvas.draw()
        # plt.tight_layout(left=0.05, right=0.95, bottom=0.1, top=0.9)  # 调整布局，但保留底部10%和顶部5%的空间给按钮
        fig.subplots_adjust(left=0.05, right=0.95, bottom=0.1, top=0.9)
        QtWidgets.QApplication.processEvents()
        plt.show(block=True)  # 确保阻塞式显示

    def updateGridLayout(self):
        """重新计算并设置每行的缩略图数量，确保缩略图的布局始终适应当前窗口大小"""
        width = self.scrollAreaThumbnails.viewport().width()
        num_columns = max(1, width // (self.thumbnail_width + 10))  # 缩略图间隔10像素   新的现在可以去展示的列数
        current_columns = self.gridLayoutThumbnails.columnCount()  # 当前已显示的列数

        if current_columns != num_columns:
            old_widgets = [self.gridLayoutThumbnails.itemAt(i).widget() for i in
                           range(self.gridLayoutThumbnails.count())]
            # 清空并重新添加所有缩略图以更新布局，需要再次使用！
            # ！！！self.clearLayout(self.gridLayoutThumbnails)会导致小部件在被重新添加到布局之前被销毁，导致old_widgets为空不显示
            for widget in old_widgets:
                self.gridLayoutThumbnails.removeWidget(widget)  # 这里只能是暂时移除
                widget.setParent(None)
            for i, widget in enumerate(old_widgets):
                row = i // num_columns
                col = i % num_columns
                self.gridLayoutThumbnails.addWidget(widget, row, col)

    def clearLayout(self, layout):
        while layout.count():
            item = layout.takeAt(0)
            widget = item.widget()
            if widget is not None:
                widget.deleteLater()

    def startColmap(self):
        # 检查目录中是否有jpg图像文件
        image_dir = os.path.join(self.WORK_DIR, "input")
        jpg_files = [f for f in os.listdir(image_dir) if f.lower().endswith('.jpg')]
        if not jpg_files:
            QtWidgets.QMessageBox.warning(self, "警告", "需要导入数据")
            return  # 如果没有jpg文件，直接返回，不执行线程

        if self.ColmapThread is None or not self.ColmapThread.isRunning():  # 进行检查和必要时重新创建。这确保了无论之前的线程是否已经完成，都可以安全地启动新的线程
            self.textEditConsole.clear()  # 清空 textEditConsole
            self.ColmapThread = ColmapWorker(self.WORK_DIR)  # 创建一个新的线程实例
            self.ColmapThread.finished.connect(self.onColmap_finished)  # 重建完成发出完成信号，然后展示结果
            self.ColmapThread.log_message.connect(self.textEditConsole.append)
            self.ColmapThread.start()
            self.tabWidget.setCurrentIndex(2)  # 跳转到标签页3
        else:
            QtWidgets.QMessageBox.warning(self, "警告", "正在构建，请稍后")

    def setPlane(self, have_plane, plane, best_image_id, best_image_name, best_points):
        """设置平面参数"""
        if isinstance(plane, list) and len(plane) == 4 and have_plane is True:
            self.have_plane = have_plane
            self.plane = np.array(plane, dtype=np.float64)
            print(f"平面参数已设置为: {self.plane}")
            self.best_image_id = best_image_id
            self.best_image_name = best_image_name
            self.best_points = best_points
        else:
            raise ValueError("平面参数必须是一个包含4个元素的列表或数组")

    def onColmap_finished(self, filepath):
        self.openPointCloud(filepath)  # 打开构建完成的点云文件
        # 得到平面参数
        # self.SemanticSegmentation_Worker = SemanticSegmentation_Worker(os.path.join(self.WORK_DIR, 'input'),
        #                                                                self.WORK_DIR, self.best_image_id,
        #                                                                self.best_image_name)
        self.SemanticSegmentation_Worker = SemanticSegmentation_Worker(self.image_paths, self.WORK_DIR)
        self.SemanticSegmentation_Worker.finished.connect(self.setPlane)
        self.SemanticSegmentation_Worker.updateView.connect(self.showImages)
        self.SemanticSegmentation_Worker.start()
        # 设置计时器检查 have_plane 状态
        self.wait_plane_timer = QtCore.QTimer(self)
        self.wait_plane_timer.timeout.connect(self.check_plane_status)
        self.wait_plane_timer.start(1000)  # 每1秒检查一次
        self.textEditConsole.append("正在等待平面检测完成...")

    def check_plane_status(self):
        if self.have_plane:
            self.wait_plane_timer.stop()
            self.textEditConsole.append("平面检测完成，开始校正对齐...")
            # 平面检测完成后，开始车牌对齐处理
            self.Align_according_to_LicensePlate_Worker = Align_according_to_LicensePlate_Worker(
                self.WORK_DIR,
                self.plane, self.best_image_name, self.best_points)
            self.Align_according_to_LicensePlate_Worker.finished.connect(self.calculating_the_scale)
            self.Align_according_to_LicensePlate_Worker.start()

    def calculating_the_scale(self, real_distances):
        """计算比例"""
        if not real_distances:
            QtWidgets.QMessageBox.warning(self, "警告", "没有检测到车牌，请检查图像数据")
            return
        self.Unit_distance_length = np.float64(self.LicensePlate) / np.float64(real_distances)  # X （米/每单位）
        print(f"根据车牌算出单位距离长度为{self.Unit_distance_length:.8f} 米/单位")

    def Projection(self, point):
        """将点投影到平面上"""
        a, b, c, d = self.plane
        # 平面方程: ax + by + cz + d = 0
        x, y, z = point
        normal = np.array([a, b, c])
        point = np.array([x, y, z])
        # 计算从点到平面的有符号距离
        distance = (a * x + b * y + c * z + d) / (a ** 2 + b ** 2 + c ** 2)
        # 用距离乘法向量，得到从点垂直指向平面的矢量
        projection_point = point - distance * normal
        return projection_point

    def export_docx(self):
        image_folder = self.save_path
        # 遍历图像并调用工具函数
        image_files = glob(os.path.join(image_folder, "*.jpg")) + glob(os.path.join(image_folder, "*.png"))
        if len(image_files) <= 0:
            QtWidgets.QMessageBox.warning(self, "警告", "没有标记过的图像文件可供生成报告")
            return
        input_path = "事故报告_模板.docx"
        output_path = "事故报告.docx"
        current_time = datetime.now().strftime("%Y年%m月%d日%H时%M分%S秒")
        address = "null"
        # 打开文档检查表格结构
        from docx import Document
        doc = Document(input_path)
        if not doc.tables or len(doc.tables) == 0:
            QtWidgets.QMessageBox.warning(self, "错误", "模板文档中没有表格")
            return
        table = doc.tables[0]

        # 计算需要的总行数
        required_rows = 3 + len(image_files)  # 头部信息占2行，从第3行开始插入图片
        current_rows = len(table.rows)
        # 如果需要，添加更多的行
        for _ in range(max(0, required_rows - current_rows)):
            table.add_row()

        fill_data = [
            {"table_index": 0, "row": 0, "col": 1, "text": current_time},
            {"table_index": 0, "row": 1, "col": 1, "text": address},
        ]
        # 从第3行开始依次插入图像
        for i, img_path in enumerate(image_files):
            fill_data.append({
                "table_index": 0,
                "row": 2 + i,  # 从第3行开始插入图
                "col": 0,
                "image_path": img_path
            })
            print(f"[+] 正在处理图像：{img_path}")
        # 保存修改后的文档以供 fill_word_template 使用
        doc.save("事故报告_模板_dynamic.docx")
        fill_word_template("事故报告_模板_dynamic.docx", output_path, fill_data)
        QtWidgets.QMessageBox.information(self, "生成报告完成",
                                          f"报告已保存到: {os.path.join(os.path.dirname(__file__), output_path)}")
        # 清理临时文件
        if os.path.exists("事故报告_模板_dynamic.docx"):
            os.remove("事故报告_模板_dynamic.docx")
        print(f"\n全部完成，共生成 {len(image_files)} 份报告。")

    # 在窗口大小发生变化时被触发
    def resizeEvent(self, event):
        super(MainWindow, self).resizeEvent(event)
        self.updateGridLayout()  # 去动态调整

    def closeEvent(self, event):
        # 先关闭所有的 VTK 渲染窗口和交互器
        if hasattr(self, 'vtkWidget'):
            self.vtkWidget.GetRenderWindow().Finalize()  # 结束渲染窗口
            self.vtkWidget.close()

        self.settings.setValue('recentFiles', self.recent_files_list)  # 保存最近文件列表到设置中
        super().closeEvent(event)


if __name__ == "__main__":
    app = QtWidgets.QApplication(sys.argv)

    splash = QtWidgets.QSplashScreen()
    splash.setPixmap(QtGui.QPixmap("_internal/splash.ico"))
    splash.show()
    app.processEvents()

    window = MainWindow()
    splash.finish(window)
    window.show()
    sys.exit(app.exec_())
